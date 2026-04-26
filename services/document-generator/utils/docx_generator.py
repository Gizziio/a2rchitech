from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from typing import List, Optional
import tempfile


def create_study_guide(
    topic: str,
    learning_objectives: List[str],
    key_terms: List[str],
    include_self_check: bool,
    citations: List[dict],
    length: str = "short"
) -> str:
    """
    Create a study guide in DOCX format.
    
    Args:
        topic: Study guide topic
        learning_objectives: List of learning objectives
        key_terms: List of key terms to define
        include_self_check: Whether to include self-check questions
        citations: List of citations to include
        length: Length of the guide ("short" or "medium")
        
    Returns:
        Path to the generated DOCX file
    """
    doc = Document()
    
    # Add title
    title_paragraph = doc.add_heading(level=0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(topic)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    
    # Add learning objectives section
    if learning_objectives:
        doc.add_heading('Learning Objectives', level=1)
        for obj in learning_objectives:
            doc.add_paragraph(obj, style='List Bullet')
    
    # Add key terms section
    if key_terms:
        doc.add_heading('Key Terms', level=1)
        for term in key_terms:
            term_paragraph = doc.add_paragraph()
            term_run = term_paragraph.add_run(f"{term}: ")
            term_run.font.bold = True
            # Add definition placeholder
            term_paragraph.add_run("Definition goes here...")
    
    # Add self-check questions if requested
    if include_self_check:
        doc.add_heading('Self-Check Questions', level=1)
        for i in range(5):  # Add 5 sample questions
            doc.add_paragraph(f"Q{i+1}: Sample question related to the topic...", style='List Number')
            # Add space for answer
            doc.add_paragraph("Answer: _________________________")
            doc.add_paragraph()  # Empty paragraph for spacing
    
    # Add citations section
    if citations:
        doc.add_heading('Sources', level=1)
        for citation in citations:
            citation_para = doc.add_paragraph()
            citation_run = citation_para.add_run(f"{citation['title']}")
            citation_run.font.underline = True
            citation_para.add_run(f"\n{citation['url']}")
    
    # Save to a temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return temp_file.name